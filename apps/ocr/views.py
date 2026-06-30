import io
import json
import logging
import os
import re
import tempfile

from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.views.decorators.http import require_POST

from .forms import UploadForm
from .services.ocr_engine import process_document
from .storage import DocumentData, TempStorage

logger = logging.getLogger(__name__)


def landing(request):
    """Landing page with upload form."""
    if request.method == 'POST':
        form = UploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.cleaned_data['file']
            file_data = uploaded_file.read()
            filename = uploaded_file.name

            ext = filename.lower()
            if ext.endswith('.pdf'):
                file_type = 'pdf'
            else:
                file_type = 'image'

            doc = TempStorage.create(filename, file_data, file_type)
            return redirect('process_upload', doc_uuid=doc.uuid)
    else:
        form = UploadForm()

    return render(request, 'ocr/upload.html', {'form': form})


def process_upload(request, doc_uuid):
    """Process an uploaded document (OCR / conversion)."""
    doc = TempStorage.get(doc_uuid)
    if not doc:
        return redirect('landing')

    try:
        TempStorage.update_status(doc_uuid, 'processing')
        extracted_text = process_document(doc)
        doc.text_path.write_text(extracted_text, encoding='utf-8')
        TempStorage.update_status(doc_uuid, 'done')
    except Exception as e:
        logger.error(f"OCR failed for {doc.filename}: {e}")
        TempStorage.update_status(doc_uuid, 'failed')

    return redirect('document_detail', doc_uuid=doc_uuid)


def document_detail(request, doc_uuid):
    """Editor page for a processed document."""
    doc = TempStorage.get(doc_uuid)
    if not doc:
        return redirect('landing')

    extracted_text = ''
    if doc.status == 'done':
        extracted_text = TempStorage.get_text(doc_uuid)

    return render(request, 'ocr/document_detail.html', {
        'doc': doc,
        'extracted_text': extracted_text,
    })


def download_text(request, doc_uuid):
    """Download extracted text in the requested format."""
    doc = TempStorage.get(doc_uuid)
    if not doc:
        return redirect('landing')

    if doc.status != 'done':
        return redirect('document_detail', doc_uuid=doc_uuid)

    file_format = request.GET.get('format', 'txt').lower()
    content = TempStorage.get_text(doc_uuid)
    base_filename = doc.filename.rsplit('.', 1)[0]

    if file_format == 'docx':
        if doc.is_pdf:
            try:
                from pdf2docx import Converter

                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_out:
                    output_path = temp_out.name

                cv = Converter(str(doc.file_path))
                cv.convert(output_path, start=0, end=None)
                cv.close()

                with open(output_path, 'rb') as f:
                    docx_data = f.read()

                os.remove(output_path)

                response = HttpResponse(
                    docx_data,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{base_filename}_converted.docx"'
                return response
            except Exception as e:
                logger.error(f"PDF to DOCX conversion failed: {e}")
                return redirect('document_detail', doc_uuid=doc_uuid)
        else:
            from docx import Document as DocxDocument

            doc_obj = DocxDocument()
            doc_obj.add_heading(f'OCR Extraction: {doc.filename}', 0)
            doc_obj.add_paragraph(content)

            buffer = io.BytesIO()
            doc_obj.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.docx"'
            return response

    elif file_format == 'xlsx':
        from openpyxl import Workbook

        clean_content = re.sub(r'<[^>]+>', '', content)
        clean_content = clean_content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')

        wb = Workbook()
        ws = wb.active
        ws.title = "OCR Results"

        for row_idx, line in enumerate(clean_content.splitlines(), start=1):
            ws.cell(row=row_idx, column=1, value=line)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.xlsx"'
        return response

    else:
        clean_content = re.sub(r'<[^>]+>', '', content)
        clean_content = clean_content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')

        response = HttpResponse(clean_content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.txt"'
        return response


@require_POST
def update_document_text(request, doc_uuid):
    """Update the extracted text for a document."""
    doc = TempStorage.get(doc_uuid)
    if not doc:
        return JsonResponse({'status': 'error', 'message': 'Document not found.'}, status=404)

    if doc.status != 'done':
        return JsonResponse({'status': 'error', 'message': 'Document not ready.'}, status=400)

    try:
        data = json.loads(request.body)
        new_text = data.get('text', '')
        new_text = bleach.clean(new_text, tags=[], strip=True)
        TempStorage.update_text(doc_uuid, new_text)
        return JsonResponse({'status': 'success'})
    except Exception as e:
        logger.error(f"Failed to update text: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_POST
def delete_document(request, doc_uuid):
    """Delete a document's temp files."""
    TempStorage.delete(doc_uuid)

    if request.headers.get('Accept') == 'application/json':
        return JsonResponse({'status': 'success'})
    return redirect('landing')
