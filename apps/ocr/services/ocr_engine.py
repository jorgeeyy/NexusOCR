import logging

import pytesseract
from django.conf import settings
from PIL import Image

from .preprocessing import preprocess_image
from .postprocessing import clean_text

logger = logging.getLogger(__name__)

pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD


def extract_text_from_image(image: Image.Image) -> str:
    processed = preprocess_image(image)
    text = pytesseract.image_to_string(processed, config='--psm 6 -c preserve_interword_spaces=1')
    return clean_text(text)


def extract_high_fidelity_pdf(file_path: str) -> str:
    try:
        import fitz
        doc = fitz.open(file_path)
        html_out = []
        for page in doc:
            html_out.append(page.get_text("html"))
        doc.close()
        return clean_text("\n".join(html_out))
    except Exception as e:
        logger.error(f"High-fidelity PDF extraction failed: {e}")
        return ""


def extract_scanned_pdf(file_path: str) -> str:
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(file_path, dpi=300, poppler_path=settings.POPPLER_PATH)
        page_texts = []
        for i, page_image in enumerate(images, start=1):
            text = extract_text_from_image(page_image)
            if text:
                page_texts.append(text)
        return "\n\n".join(page_texts)
    except Exception as e:
        logger.error(f"Scanned PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to convert PDF: {e}")


def extract_pdf_formatted_html(file_path: str) -> str:
    """Extract PDF text as HTML preserving alignment and basic formatting.

    Uses PyMuPDF (fitz) to read text blocks with position information.
    Detects left/right/center alignment from x-coordinates and bold/italic
    from font names, then reconstructs the layout as HTML.
    """
    import fitz
    doc = fitz.open(file_path)
    page_html_parts = []

    for page in doc:
        pw = page.rect.width
        blocks = page.get_text("dict")["blocks"]

        raw_lines = []
        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                lx0, ly0, lx1, ly1 = line["bbox"]
                if lx0 > pw * 0.6:
                    align = "right"
                elif lx0 > pw * 0.3 and lx1 < pw * 0.7:
                    align = "center"
                else:
                    align = "left"

                text_parts = []
                for span in line["spans"]:
                    raw = span["text"]
                    if not raw:
                        continue
                    txt = raw.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    font = (span.get("font") or "").lower()
                    is_bold = "bold" in font or "bd" in font or "heavy" in font or "black" in font
                    is_italic = "italic" in font or "it" in font or "oblique" in font
                    if is_bold and is_italic:
                        txt = f"<strong><em>{txt}</em></strong>"
                    elif is_bold:
                        txt = f"<strong>{txt}</strong>"
                    elif is_italic:
                        txt = f"<em>{txt}</em>"
                    text_parts.append(txt)

                raw_lines.append((ly0, align, "".join(text_parts)))

        if not raw_lines:
            continue

        raw_lines.sort(key=lambda x: x[0])

        # Group consecutive lines into paragraphs by vertical proximity.
        # Lines close together (within 20px, typical line-height at 12pt)
        # with the same alignment belong to the same paragraph.
        paragraphs = []
        cur_align = None
        cur_y = None
        cur_lines = []

        for y, align, text in raw_lines:
            if cur_lines and (align != cur_align or (cur_y is not None and y - cur_y > 20)):
                paragraphs.append((cur_align, cur_lines))
                cur_lines = []
            cur_align = align
            cur_y = y
            cur_lines.append(text)

        if cur_lines:
            paragraphs.append((cur_align, cur_lines))

        for align, lines in paragraphs:
            inner = "<br>".join(lines)
            if align == "left":
                page_html_parts.append(f"<p>{inner}</p>")
            else:
                page_html_parts.append(f'<p style="text-align:{align}">{inner}</p>')

    doc.close()
    return "\n".join(page_html_parts)


def _run_to_html(run) -> str:
    text = run.text
    if not text:
        return ""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<span style=\"text-decoration:underline\">{text}</span>"
    return text


def _paragraph_to_html(para) -> str:
    style_name = para.style.name.lower() if para.style and para.style.name else "normal"
    tag_map = {
        "heading 1": "h1", "heading 2": "h2", "heading 3": "h3",
        "heading 4": "h4", "heading 5": "h5", "heading 6": "h6",
    }
    tag = tag_map.get(style_name, "p")

    inner = "".join(_run_to_html(r) for r in para.runs)
    if not inner:
        inner = para.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    return f"<{tag}>{inner}</{tag}>"


def _table_to_html(table) -> str:
    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            cell_html = "".join(_paragraph_to_html(p) for p in cell.paragraphs)
            cells_html.append(f"<td>{cell_html}</td>")
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    return f"<table>{''.join(rows_html)}</table>"


def convert_pdf_to_docx_and_extract_html(file_path: str, docx_path: str) -> str:
    """Convert a PDF to .docx preserving layout, then extract formatted HTML.

    Returns formatted HTML on success, or empty string if the PDF has
    no extractable text (e.g. scanned pages) so the caller can fall back to OCR.
    """
    try:
        from pdf2docx import Converter
        cv = Converter(file_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    except Exception as e:
        logger.warning(f"pdf2docx conversion failed (PDF may be scanned): {e}")
        return ""

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(docx_path)
        return _docx_to_html(doc)
    except Exception as e:
        logger.warning(f"docx-to-html extraction failed: {e}")
        return ""


def _docx_to_html(doc) -> str:
    """Convert a python-docx Document to formatted HTML."""
    parts = []

    # Build ordered list of all elements (paragraphs + tables) by their XML order
    body = doc.element.body
    elements = []

    import lxml.etree as ET
    para_elems = {p._element: p for p in doc.paragraphs}
    table_elems = {t._element: t for t in doc.tables}

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p' and child in para_elems:
            elements.append(('p', para_elems[child]))
        elif tag == 'tbl' and child in table_elems:
            elements.append(('tbl', table_elems[child]))

    for elem_type, elem in elements:
        if elem_type == 'p':
            html = _paragraph_to_html(elem)
            if html:
                parts.append(html)
        elif elem_type == 'tbl':
            parts.append(_table_to_html(elem))

    return "\n".join(parts)


def process_document(doc) -> str:
    """Process a DocumentData instance and extract text.

    For PDFs: extracts formatted HTML via PyMuPDF (preserving alignment,
    bold/italic). Also saves a layout-preserving .docx via pdf2docx
    for export. Falls back to OCR if the PDF has no selectable text.
    For images: uses Tesseract OCR with space preservation.
    """
    from ..storage import TempStorage

    file_path = str(doc.file_path)
    TempStorage.update_status(doc.uuid, 'processing')

    try:
        if doc.is_pdf:
            text = extract_pdf_formatted_html(file_path)
            if len(text.strip()) < 100:
                logger.info(f"PDF seems scanned, switching to OCR for {doc.filename}")
                text = extract_scanned_pdf(file_path)

            # Save layout-preserving .docx for export (best effort)
            try:
                from pdf2docx import Converter
                cv = Converter(file_path)
                cv.convert(str(doc.docx_path), start=0, end=None)
                cv.close()
            except Exception as e:
                logger.warning(f"Could not create export .docx for {doc.filename}: {e}")

        elif doc.is_image:
            image = Image.open(file_path)
            text = extract_text_from_image(image)

        else:
            raise ValueError(f"Unsupported file type: {doc.file_extension}")

        TempStorage.update_status(doc.uuid, 'done')
        return text

    except Exception as e:
        TempStorage.update_status(doc.uuid, 'failed')
        logger.error(f"OCR failed for {doc.filename}: {e}")
        raise
